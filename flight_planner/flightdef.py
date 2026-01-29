"""
Class definitions for WayPoint and FlightDef objects, plus some utility
functions.
"""

import os
import numpy as np
import cartopy.geodesic as cgeodesic
from docx import Document, shared
import gpxpy, gpxpy.gpx

from .user_config import airports # Needed to lock WPs to nearby airports
from .user_config import aircrafts, legtype_spds # Needed for setting spds at load

# Set default directory fddir for saving flight defs, allowing for frozen build

class WayPoint(object):
    """A single position.

    Initialise with:
        lon, lat (float)        : Position in decimal degrees
        alt (float, default 0.0): Altitude in ft
        name (str, default 'NONAME'): A single word (using A-Z and 0-9 only)
        desc (str, default '""')    : A short description (<40 chars)
        legtype (str, default 'transit'): Used for computing times, allowed 
             legtypes are defined by the legtype_spds dict in user_config file

    (the requirements on name and desc are for compatability with foreflight)
    """
    def __init__(self, lon, lat, alt=0.0, name='NONAME', desc='',
                 legtype='transit'):
        self.lon = lon
        self.lat = lat
        self.alt = alt
        self.name = name
        self.desc = desc
        self.legtype = legtype

    def __repr__(self):
        # Neat data dump for writing save file
        s = '{0:6}, {2: 7.3f}, {3: 6.3f}, {4:5.0f}, {5}, {1}'.\
            format(self.name[:6], self.desc, self.lon, self.lat, self.alt,
                   self.legtype)
        return s

    def __str__(self):
        # Human-readable output
        s = '{0}:\t{1}, {2:.0f}ft [{4}] ({3})'.\
            format(self.name, decdeg2intdegdecmin_string(self.lon, self.lat),
                   self.alt, self.desc, self.legtype)
        return s

    def print_foreflight(self):
        # Format for CSV foreflight file
        s = '{0},{1},{2:.3f},{3:.3f}'.\
            format(self.name, self.desc if self.desc != '' else '""', self.lat,
                   self.lon)
        return s

    def pt(self):
        return self.lon, self.lat


def WayPoint_from_repr(s):
    """Helper function to recreate a WayPoint from its __repr__"""
    vals = [a.strip() for a in s.split(',')]
    return WayPoint(float(vals[1]), float(vals[2]), alt=float(vals[3]),
                    name=vals[0], desc=vals[5], legtype=vals[4])

    
def WayPoint_from_airport(airport):
    """Helper function to create WayPoint for a given airport code"""
    a = airports[airport]
    return WayPoint(a[0], a[1], alt=a[2], name=airport)


def is_near_airport(waypoint, wplock=100):
    for airport, pt in airports.items():
        dist = greatcircle((waypoint.lon, waypoint.lat), (pt[0], pt[1]))
        if dist < wplock: return airport
    return None


def is_near_other_waypoints(waypoint, others, wplock=100):
    for i, other in enumerate(others):
        dist = greatcircle((waypoint.lon, waypoint.lat),
                           (other.lon, other.lat))
        if dist < wplock: return i
    return None
            

def lock_WayPoint(waypoint, wplock=100):
    """Lock WayPoint to nearby airport within wplock km"""
    near_airport = is_near_airport(waypoint, wplock=wplock)
    if near_airport: return WayPoint_from_airport(near_airport)
    return waypoint


def locked_WayPoint(lon, lat, alt=0.0, name='NO_NAME', desc='',
                    legtype='transit', wplock=100):
    """Create a WayPoint but locks to airports within wplock km"""
    waypoint = WayPoint(lon, lat, alt=alt, name=name, desc=desc,
                        legtype=legtype)
    return lock_WayPoint(waypoint, wplock=wplock)


class FlightDef(object):
    """A flight definition consisting of a list of waypoints.

    Initialise with:
        waypoints (list): A list of WayPoint objects
        name (str)      : Name of flight (used for save filenames)
        aircraft (str, optional): Name of aircraft
        legtype_spds (dict, optional): speeds used for each leg type

    Note: To modify the flight definition, just use the list methods (append,
    insert, delete etc) on FlightDef.waypoints

    Main methods:
        lons, lats, alts: Return array of coordinate values
        total_summary   : Return stats for full flight
        leg_summary(i)  : Return stats for leg i
    """

    def __init__(self, waypoints=[], name='no_name', aircraft=None,
                 legtype_spds={'transit': 100}, datapath=os.getcwd()):
        self.waypoints = waypoints
        self.name = name
        self.aircraft = aircraft
        self.legtype_spds = legtype_spds
        self.set_fddir(datapath)

    def set_fddir(self, datapath):
        self.fddir = os.path.join(datapath, 'data/flight_defs')
        if os.path.exists(self.fddir) is False:
            print('\nCreating flight def directory:\n{}'.format(self.fddir))
            os.makedirs(self.fddir, exist_ok=True)
        else:
            print('Setting flight def directory: {}'.format(self.fddir))

    def __repr__(self):
        """Data dump for writing save file"""
        s = '{}'.format(self.name)
        for waypoint in self.waypoints: s += '\n' + waypoint.__repr__()
        return s

    def __str__(self):
        """Human-readable output"""
        s = 'Name: {}'.format(self.name)
        s += '\nAircraft: {}'.format(self.aircraft)
        s += '\nSummary: ' + self.total_summary()
        s += '\nWaypoints:'
        for i, waypoint in enumerate(self.waypoints):
            s += '\n  {}'.format(waypoint.__str__())
        s += '\nLegs:'
        for i, leg in enumerate(self.legs()):
            s += '\n  {3}: {0:.2f}nm, {1} @ {2:.0f}kts'.\
                format(km2nm(self.leg_dist(leg)),
                       hr2str(self.leg_time(leg)),
                       self.legtype_spds[leg[0].legtype],
                       i)
        return s

    def print_gpx(self):
        """GPX output for importing into Windy"""
        gpx = gpxpy.gpx.GPX()
        # Create empty GPX track:
        gpx_track = gpxpy.gpx.GPXTrack()
        gpx.tracks.append(gpx_track)
        # Add segment in our GPX track:
        gpx_segment = gpxpy.gpx.GPXTrackSegment()
        gpx_track.segments.append(gpx_segment)
        # Add waypoints points:
        for wp in self.waypoints:
            gpx_segment.points.append(
                gpxpy.gpx.GPXTrackPoint(wp.lat, wp.lon,
                                        elevation=ft2m(wp.alt)))
        return gpx.to_xml()

    def print_speeds(self):
        """Print list of legtype and their speeds"""
        return ', '.join(['{}: {}kts'.format(k, v)
                          for k, v in self.legtype_spds.items()])

    def lons(self):
        """Get an array of longitude values"""
        return np.array([waypoint.lon for waypoint in self.waypoints])

    def lats(self):
        """Get an array of latitude values"""
        return np.array([waypoint.lat for waypoint in self.waypoints])

    def alts(self):
        """Get an array of altitude values"""
        return np.array([waypoint.alt for waypoint in self.waypoints])

    def legs(self):
        """Return a list of WayPoint pairs defining each leg"""
        return [(self.waypoints[i], self.waypoints[i+1])
                for i in range(len(self.waypoints) - 1)]

    # Calculations using track definition
    def total_summary(self):
        return '{0:.1f}nm, {1} '.\
            format(km2nm(self.total_dist()), hr2str(self.total_time())) +\
                '(' + self.print_speeds() + ')'

    def total_dist(self, cumulative=False):
        dists = [self.leg_dist(leg) for leg in self.legs()]
        if cumulative: return np.cumsum(dists)
        return np.sum(dists)

    def total_time(self, cumulative=False):
        times = [self.leg_time(leg) for leg in self.legs()]
        if cumulative: return np.cumsum(times)
        return np.sum(times)

    def leg_summary(self, i):
        leg = self.legs()[i]
        return 'Leg {0}: {2:.1f}nm, {3} [{1}]'.\
            format(i+1, leg[0].legtype,
                   km2nm(self.leg_dist(leg)), hr2str(self.leg_time(leg)))

    def leg_dist(self, leg):
        return greatcircle(leg[0].pt(), leg[1].pt())

    def leg_time(self, leg):
        time = km2nm(self.leg_dist(leg)) / self.legtype_spds[leg[0].legtype]  # hr
        return time
    
    def leg_segments(self):
        return [((leg[0].lon, leg[0].lat), (leg[1].lon, leg[1].lat))
                for leg in self.legs()]
    
    # Save functions
    def savedat(self, fn=None):
        """Save FlightDef as txt file for reloading later with loaddat"""
        if fn is None: fn = os.path.join(self.fddir, self.name+'.dat')
        print('\nSaving flight def to file: {}'.format(fn))
        with open(fn, 'w') as f:
            print(self.__repr__())
            f.write(self.__repr__())
        
    def savetxt(self, fn=None):
        """Save FlightDef as txt file that is nice to read"""
        if fn is None: fn = os.path.join(self.fddir, self.name+'.txt')
        print('\nExporting readable flight def to file: {}'.format(fn))
        with open(fn, 'w') as f:
            print(self.__str__())
            f.write(self.__str__())
            
    def savegpx(self, fn=None):
        """Save FlightDef as gpx file that can be imported to windy"""
        if fn is None: fn = os.path.join(self.fddir, self.name+'.gpx')
        print('\nExporting GPX flight def file: {}'.format(fn))
        with open(fn, 'w') as f:
            print(self.print_gpx())
            f.write(self.print_gpx())
        
    def savecsv(self, fn=None):
        """Save FlightDef as CSV file formatted for foreflight"""
        if fn is None: fn = os.path.join(self.fddir, self.name+'.csv')
        print('\nExporting ForeFlight file: {}'.format(fn))
        with open(fn, 'w') as f:
            print('WAYPOINT,DESCRIPTION,LATITUDE,LONGITUDE')
            f.write('WAYPOINT,DESCRIPTION,LATITUDE,LONGITUDE\n')
            for waypoint in self.waypoints:
                print(waypoint.print_foreflight())
                f.write(waypoint.print_foreflight()+'\n')
            
    def savedoc(self, fn=None, fnfig=None):
        """Create Word document for flight sortie"""
        if fn is None: fn = os.path.join(self.fddir, self.name+'.docx')
        f = Document()
        style = f.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = shared.Pt(12)
        font.bold = False
        f.add_paragraph().add_run(f'Flight code and title: {self.name}'
                                  '\n\nMission Scientist(s):'
                                  '\n\nObjectives:'
                                  '\n\nFlight Details:').font.bold=True
        rows = self._create_wptable()        
        # add bullet points for details
        base = rows[0][0]
        end = rows[-1][0]
        f.add_paragraph(f'{base} to {end}', style='List Bullet')
        f.add_paragraph(f'Estimated flight time = {hr2str(self.total_time())}',
                        style='List Bullet')      
        dist = round(km2nm(self.total_dist()), 1)
        f.add_paragraph(f'Total distance = {dist} nm', style='List Bullet')
        f.add_paragraph('Take off time = ', style='List Bullet')
        f.add_paragraph().add_run('Weather Concerns:'
                                  '\n\nSafety:'
                                  '\n\nDiagram:').font.bold=True
        # Save current fig for inclusion in word doc
        if fnfig: f.add_picture(fnfig, width=shared.Inches(5))
        f.add_paragraph().add_run('Waypoints:').font.bold=True
        # create table for wps
        t = f.add_table(rows=len(rows)+1, cols=7)
        headings = t.rows[0].cells
        headings[0].text = 'Code'
        headings[1].text = 'Lon'
        headings[2].text = 'Lat'
        headings[3].text = 'Alt'
        headings[4].text = 'Leg Type'
        headings[5].text = 'Description'  
        headings[6].text = 'Elapsed Time'                 
        for i in range(len(rows)):
            row = t.rows[i+1].cells
            for j in range(7):
                row[j].text = str(rows[i][j])
        t.style = 'Table Grid'
        print('\nExporting flight sortie to file: {}'.format(fn))
        f.save(fn)
    
    def _create_wptable(self):
        rows = self.__repr__().split('\n')
        rows.pop(0)
        elapsed_time = 0
        n = len(rows)
        for i, row in enumerate(rows):
            row = [cell.strip() for cell in row.split(',')]
            # convert to decimal minutes
            decmin = decdeg2intdegdecmin_string(float(row[1]), float(row[2]))
            row[1], row[2] = decmin.split()
            # add elapsed time
            row.append(hr2str(elapsed_time))
            rows[i] = row
            if i < n-1: elapsed_time += self.leg_time(self.legs()[i])
        return rows
    
    
# Load functions
def loaddat(fn):
    """Load FlightDef from dat file"""
    print('\nLoading flight def file: {}'.format(fn))
    with open(fn, 'r') as f:
        name = f.readline().split('\n')[0]
        waypoints = []
        for line in f.readlines():
            next_waypoint = WayPoint_from_repr(line.split('\n')[0])
            waypoints.append(next_waypoint)
    datapath = os.path.dirname(fn).split('data/flight_defs')[0]
    new_flightdef = FlightDef(waypoints=waypoints, name=name,
                              aircraft=aircrafts[0],
                              legtype_spds=legtype_spds[aircrafts[0]],
                              datapath=datapath)
    print(new_flightdef)
    return new_flightdef

def loadgpx(fn):
    """Load FlightDef from gpx file"""
    print('\nLoading flight def file: {}'.format(fn))
    with open(fn, 'r') as gpx_file: gpx = gpxpy.parse(gpx_file)
    lons = [Point.longitude for Point in gpx.tracks[0].segments[0].points]
    lats = [Point.latitude for Point in gpx.tracks[0].segments[0].points]
    waypoints = [WayPoint(lon, lat) for lon, lat in zip(lons, lats)]
    datapath = os.path.dirname(fn).split('data/flight_defs')[0]
    new_flightdef = FlightDef(waypoints=waypoints, name='MMDDa',
                              aircraft=aircrafts[0],
                              legtype_spds=legtype_spds[aircrafts[0]],
                              datapath=datapath)
    print(new_flightdef)
    return new_flightdef


# Convert functions
def decdeg2intdegdecmin(lon, lat):
    """Convert lon and lat in decimal degrees to (deg [int], decimal minute, H).
    where H is one of 'E', 'W' (lon) and 'N', 'S' (lat).
    """
    lonH = 'E' if lon >= 0 else 'W'
    latH = 'N' if lat >= 0 else 'S'
    abslon = np.abs(lon)
    abslat = np.abs(lat)
    intlon = int(np.floor(abslon))
    intlat = int(np.floor(abslat))
    minlon = (abslon - intlon) * 60
    minlat = (abslat - intlat) * 60
    return [intlon, minlon, lonH, intlat, minlat, latH]

def decdeg2intdegdecmin_string(lon, lat):
    return u'{0:03d}\N{DEGREE SIGN}{1:05.2f}{2} '\
        '{3:02d}\N{DEGREE SIGN}{4:05.2f}{5}'.\
        format(*decdeg2intdegdecmin(lon, lat))

def intdegdecminstr2lonlat(d):
    """Convert single string of form 820000N0000000E, where numbers are 
    deg and decimal minutes, to (lon, lat) in degrees.

    (e.g. FIR boundaries coordinates from Dan Breedon)
    """
    intlat = int(d[0:2])
    intminlat = int(d[2:4])
    decminlat = int(d[4:6])
    latH = d[6]
    lat = intlat + (intminlat + decminlat / 100.) / 60.
    if latH == 'N':
        pass
    elif latH == 'S':
        lat = -lat
    else:
        print('ERROR: Latitude should be N or S')
    intlon = int(d[7:10])
    intminlon = int(d[10:12])
    decminlon = int(d[12:14])
    lonH = d[14]
    lon = intlon + (intminlon + decminlon / 100.) / 60.
    if lonH == 'E':
        pass
    elif lonH == 'W':
        lon = -lon
    else:
        print('ERROR: Longitude should be E or W')
    return lon, lat

def km2nm(dist):
    # Convert km to nm
    return dist / 1.852

def nm2km(dist):
    # Convert nm to km
    return dist * 1.852

def ft2m(alt):
    # Convert ft to m
    return alt / 0.3048

def m2ft(alt):
    # Convert m to ft
    return alt * 0.3048

def hr2str(hrs):
    # Convert number of hours to a string XhrYZ
    hri = int(np.floor(hrs))
    mini = int(np.round((hrs - hri) * 60))
    return '{0}hr{1}min'.format(hri, mini)
    
def greatcircle(pt0, pt1):
    """Return great circle distance between a pair of (lon, lat) points in km
    """
    geodesic = cgeodesic.Geodesic()
    dist = geodesic.inverse(pt0, pt1)[0, 0] * 1e-3 # m to km
    return dist